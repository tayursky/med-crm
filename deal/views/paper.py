from decimal import Decimal
import json
import re
from io import StringIO, BytesIO
from datetime import date, datetime, timedelta

from django.db.models import Q
from django.http import HttpResponse, JsonResponse
from django.views.generic.edit import View
from django.shortcuts import render, render_to_response
from django.conf import settings
from django.utils.http import urlquote
from django.template.loader import render_to_string

from docx import Document
from docx.shared import Inches

# from weasyprint import HTML

from absolutum.settings import BASE_DIR
from deal.models import Client, Deal, DealPerson, Stage, Service, ServiceGroup
from identity.models import Person
from utils.date_time import get_month_name


class DealDocView(View):
    model = Deal
    doc_type = None
    deal = None
    person = None
    permissions = []
    template_name = 'paper/test.jinja2'

    def dispatch(self, request, *args, **kwargs):
        self.doc_type = request.GET.get('doc_type')
        self.permissions = Deal.get_permissions(self.request)
        try:
            self.deal = self.model.objects.get(pk=request.GET.get('deal'))
        except (ValueError, Deal.DoesNotExist):
            pass
        try:
            self.person = Person.objects.get(pk=request.GET.get('person'))
        except (ValueError, Client.DoesNotExist):
            pass

        return super().dispatch(request, *args, **kwargs)

    def get_data(self, *args, **kwargs):
        _today = date.today()
        age = (_today - self.person.birthday).days / 365 if self.person.birthday else None
        kid = ''
        if age and (age < 18 and self.person.family.all().exists()):
            kid = self.person
            self.person = self.person.family.first().relative

        answer = dict(
            template_file=BASE_DIR + '/other/docx/',
            data=dict(
                title=self.deal.cache.get('title', None) if self.deal else 'Договор',
                id=self.deal.id,
                deal_cost=str(self.deal.cost)[:-3],
                deal_cost_bone=str(self.deal.cost)[-2:],
                deal_city=self.deal.branch.city.name,
                kid='',
                full_name=self.person.cache.get('full_name'),
                phone=self.person.cache.get('phone', ''),
                email=self.person.get_email(),
                birth_dd=self.person.birthday.strftime('%d') if self.person.birthday else ' ',
                birth_mm=self.person.birthday.strftime('%m') if self.person.birthday else ' ',
                birth_yy=self.person.birthday.strftime('%Y') if self.person.birthday else ' ',
                dd=_today.strftime('%d'),
                mm=get_month_name(_today.strftime('%m')),
                yy=_today.strftime('%Y'),
                address=self.person.address or '',
                passport_number=self.person.passport_number or '',
                passport_issued=self.person.passport_issued or '',
            )
        )
        if kid:
            answer['data'].update(dict(
                kid=kid,
                kid_yy=kid.birthday.strftime('%Y') if kid.birthday else ' ',
                kid_address=kid.address or '',
                kid_p_number=kid.passport_number or '',
                kid_p_issued=kid.passport_issued or '',
            ))

        if self.doc_type == 'contract':
            answer['template_file'] += 'contract_massage.docx'
            answer['file_name'] = 'Договор_%s' % answer['data'].get('full_name')
            answer['data'].update(dict(

            ))
        elif self.doc_type == 'consent_personal':
            if kid:
                answer['template_file'] += 'personal_data_kid.docx'
                answer['file_name'] = 'Согласие_%s' % answer['data'].get('kid')
            else:
                answer['template_file'] += 'personal_data.docx'
                answer['file_name'] = 'Согласие_%s' % answer['data'].get('full_name')
            answer['data'].update(dict(

            ))
        answer['file_name'].replace(' ', '_')
        return answer

    def get(self, request, *args, **kwargs):
        if 'view' not in self.permissions:
            return JsonResponse(dict(answer='No permissions'), safe=False)
        if self.request.GET.get('debug'):
            return render_to_response('debug.jinja2', locals())

        _data = self.get_data()
        _doc = Document(_data['template_file'])

        for regex, replace in _data['data'].items():
            print(regex, replace)
            _regex = re.compile(r'{{ %s }}' % regex)
            _replace = r'%s' % replace
            doc_replace_regex(_doc, _regex, _replace)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=%s.docx' % urlquote(_data['file_name'])
        _doc.save(response)
        return response

        # return JsonResponse(context, safe=False)


def doc_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                doc_replace_regex(cell, regex, replace)


class DealPaperView(View):
    model = Deal
    doc_type = None
    deal = None
    person = None
    permissions = []
    template_name = 'paper/test.jinja2'

    def dispatch(self, request, *args, **kwargs):
        self.doc_type = kwargs.get('doc_type')
        self.permissions = Deal.get_permissions(self.request)
        try:
            self.deal = self.model.objects.get(pk=request.GET.get('deal'))
        except (Deal.DoesNotExist, ValueError):
            pass
        try:
            self.person = Person.objects.get(pk=request.GET.get('person'))
        except (ValueError, Client.DoesNotExist):
            pass
        return super().dispatch(request, *args, **kwargs)

    def get(self, request, *args, **kwargs):
        if 'view' not in self.permissions:
            return JsonResponse(dict(answer='No permissions'), safe=False)

        context = dict(
            title=self.deal.cache.get('title', None) if self.deal else 'Договор',
            deal=self.deal,
            person=self.person,
        )

        if self.request.GET.get('debug'):
            return render_to_response('debug.jinja2', locals())

        return render_to_response('debug.jinja2', locals())

        # html = render_to_string(self.template_name, context, request)
        # uri = request.build_absolute_uri()
        # pdf = HTML(string=html, base_url=uri, encoding="UTF-8").write_pdf(presentational_hints=True)
        # response = HttpResponse(pdf, content_type='application/pdf')
        # response['Content-Disposition'] = 'inline; filename=%s.pdf' % urlquote(context['title']).lower()
        # return response
